#coding=utf-8
#version: 7.0.1
import json
import sys
import os
import re
import imghdr
import tkinter as tk
from tkinter import messagebox
from time import sleep

import docx
import requests
from bs4 import BeautifulSoup
from docx.shared import Inches
from PIL import Image
from PIL import ImageTk


'''
下載 Dcard 文章以及圖片
'''


class Dcard():
    def __init__(self, url):
        self.__url = url
        self.dcard_text = []
        self.dcard_image_url_count = 0
        self.dcard_sentence_count = 0
        self.dcard_title = ''
        self.dcard_fail_img_list = {}
        self.dcard_headers = {
            'user-agent':
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 11_4) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36 Edg/91.0.864.59'
        }

    # 分析dcard文章網址
    def dcard_analysis(self):

        # 取得文章編號
        post_num = self.__url.rsplit('/', 1)[1]

        # 利用官方 API 取得文章內容
        dcard_api_url = 'https://www.dcard.tw/_api/posts/' + post_num

        r = requests.get(dcard_api_url, headers=self.dcard_headers)

        if r.status_code == requests.codes.ok:

            dict_json = r.json()

            # 抓取文章標題
            self.dcard_title = dict_json['title'] 

            # 取得作者名稱
            if dict_json['anonymousSchool'] == True:
                self.dcard_title += '-匿名' 
            else:
                self.dcard_title += '-' + dict_json['school']

            # 檢查有無 / 存在
            if '\\' in self.dcard_title:
                tmp = []
                for i in self.dcard_title.split('\\'):
                    tmp.append(i)
                self.dcard_title = '+'.join(tmp)
            elif '/' in self.dcard_title:
                tmp = []
                for i in self.dcard_title.split('/'):
                    tmp.append(i)
                self.dcard_title = '+'.join(tmp)

            # 創建資料夾
            try:
                mkdir(self.dcard_title)
            except OSError as e:
                print(e)
                print('Filename have some problem, You should change by yourself')
                self.dcard_title = 'tmp_dcard_file_name'
                mkdir(self.dcard_title)

            # 抓取文章內容及圖片網址
            list_text = dict_json['content'].split('\n')
            for text in list_text:
                if 'http' in text:
                    self.dcard_image_url_count += 1
                    try:
                        self.dcard_image_download(text)
                    except Exception as e:
                        print(e)
                        self.dcard_fail_img_list[str(self.dcard_image_url_count)] = text
                        text_update('下載圖片過程發生問題\n')
                    self.dcard_text.append(text)
                else:
                    self.dcard_text.append(text)
                self.dcard_sentence_count += 1

            if self.dcard_image_url_count > 9:
                text_update('文章總共 %s 張圖片\n' % str(self.dcard_image_url_count))
            else:
                text_update('文章總共  %s 張圖片\n' % str(self.dcard_image_url_count))
            text_update('文章圖片下載完成\n')

            # 取得留言數量
            commentcnt = dict_json['commentCount']
            total_cnt = (commentcnt / 30) + 1

            # API 一次是 30 筆，所以要回圈取得資訊
            comments_img_cnt = 0
            cycle_cnt = 0
            error_cnt = 0
            while cycle_cnt < total_cnt:

                # 利用官方 API 取得文章留言
                if cycle_cnt == 0:
                    comment_api_url = 'https://www.dcard.tw/_api/posts/' + post_num + '/comments'
                else:
                    comment_api_url = 'https://www.dcard.tw/_api/posts/' + post_num + '/comments' + '?after=' + str(cycle_cnt * 30)

                r2 = requests.get(comment_api_url, headers=self.dcard_headers)

                if r2.status_code == requests.codes.ok:

                    comment_json = r2.json()

                    # 抓取文章留言及圖片網址
                    for comments in comment_json:
                        if comments['hiddenByAuthor'] == True or comments['hidden'] == True:
                            continue
                        comments_time_hour = re.search('(.*)T(.*):(.*):(.*)', comments['updatedAt']).group(2)
                        comments_time_minute = re.search('(.*)T(.*):(.*):(.*)', comments['updatedAt']).group(3)
                        comments_time = str(int(comments_time_hour) + 8) + ':' + comments_time_minute
                        if comments['host'] == True:
                            comments_info = 'B' + str(comments['floor']) + ' - 原PO - ' + comments_time
                        else:
                            comments_info = 'B' + str(comments['floor']) + ' - ' + comments['school'] + ' - ' + comments_time
                        self.dcard_text.append('-------------------------')
                        self.dcard_text.append(comments_info)
                        self.dcard_sentence_count += 2
                        list_comment = comments['content'].split('\n')
                        for text in list_comment:
                            if 'http' in text:
                                self.dcard_image_url_count += 1
                                comments_img_cnt += 1
                                try:
                                    self.dcard_image_download(text)
                                except Exception as e:
                                    print(e)
                                    text_update('下載留言圖片過程發生問題\n')
                                self.dcard_text.append(text)
                            else:
                                self.dcard_text.append(text)
                            self.dcard_sentence_count += 1
                else:
                    print('Have some Error')
                    error_cnt += 1
                    if error_cnt > 100:
                        print('Too much Error, so only save article and pic')
                        break
                    sleep(0.1)
                    continue
                # 完成一圈加一
                cycle_cnt += 1
            if comments_img_cnt > 9:
                text_update('留言總共 %s 張圖片\n' % str(comments_img_cnt))
            else:
                text_update('留言總共  %s 張圖片\n' % str(comments_img_cnt))
            text_update('留言圖片下載完成\n')

            # 下載 Dcard 文章
            try:
                self.dcard_txt_download(self.__url)
            except Exception as e:
                print(e)
                text_update('下載文章過程發生問題\n')

            # 輸出失敗資訊
            for key in self.dcard_fail_img_list:
                print(key, '->', self.dcard_fail_img_list[key])
        else:
            print(r.status_code)

    # 下載圖片
    def dcard_image_download(self, dcardimageurl):

        # 過濾掉跟網址黏在一起的字
        dcardimageurl = 'http' + re.search('(.*)http(.*)', dcardimageurl).group(2)

        # 將網址換成正確的格式
        if 'imgur.dcard.tw' in dcardimageurl:
            dcardimageurl = dcardimageurl.replace('imgur.dcard.tw',
                                                  'i.imgur.com')
        elif 'megapx-assets.dcard.tw' in dcardimageurl:
            dcardimageurl_new = dcardimageurl.rsplit('/', 1)[0] + '/full.jpeg'
        elif 'i.imgur.com' in dcardimageurl:
            dcardimageurl_new = dcardimageurl
        elif 'vivid/videos' in dcardimageurl:
            self.dcard_video_download(dcardimageurl)
            return
        else:
            self.dcard_image_url_count -= 1
            return

        filename = str(self.dcard_image_url_count) + '.jpg'
        img = requests.get(
            dcardimageurl_new, headers=self.dcard_headers)
        imgcontent = img.content
        if img.status_code != requests.codes.ok:
            imgcontent = requests.get(
                dcardimageurl, headers=self.dcard_headers).content
        with open(self.dcard_title + '/' + filename, 'wb') as code:
            code.write(imgcontent)
            if self.dcard_image_url_count > 9:
                text_update('第 %s 張圖片下載\n' % str(self.dcard_image_url_count))
            else:
                text_update('第  %s 張圖片下載\n' % str(self.dcard_image_url_count))

    # 下載影片
    def dcard_video_download(self, dcardimageurl):

        filename = str(self.dcard_image_url_count) + '.mov'
        try:
            r = requests.get(
                    dcardimageurl, headers=self.dcard_headers)
            soup = BeautifulSoup(r.text, 'html.parser')
            video_url = soup.select('source[src]')
            video = requests.get(
                video_url[0].get('src'), headers=self.dcard_headers)
            with open(self.dcard_title + '/' + filename, 'wb') as code:
                code.write(video.content)
                if self.dcard_image_url_count > 9:
                    text_update('第 %s 影片下載\n' % str(self.dcard_image_url_count))
                else:
                    text_update('第  %s 影片下載\n' % str(self.dcard_image_url_count))
        except Exception as e:
            print(e)

    # 下載dcard文章
    def dcard_txt_download(self, url):
        text_update('------------------------------\n')
        text_update('正在下載文章\n')
        filename = self.dcard_title + '.txt'
        with open(self.dcard_title + '/' + filename, 'a') as code:
            for i in range(self.dcard_sentence_count):
                code.write(self.dcard_text[i])
                code.write('\n')
            code.write('\n')
            code.write("文章網址 : " + url)
        text_update('文章儲存完成\n')
        text_update('------------------------------\n')


'''
下載 PTT 文章以及圖片
'''


class Ptt():
    def __init__(self, url):
        self.url = url
        self.format_data = []
        self.ptt_information = []
        self.txt_data = []
        self.ptt_tmp = []
        self.image_url_count = 0
        self.sentence_count = 0
        self.ptt_cookies = {'over18': '1'}

    def analysis(self):
        r = requests.get(self.url, cookies=self.ptt_cookies)
        if r.status_code == requests.codes.ok:

            soup = BeautifulSoup(r.text, 'html.parser')
            ptt = soup.find_all('span', class_='article-meta-value')

            # 作者: ptt_information[0], 看板: ptt_information[1], 標題: ptt_information[2]
            for i in ptt:
                self.ptt_information.append(i.text)

            if '/' in self.ptt_information[2]:
                for i in self.ptt_information[2].split('/'):
                    self.ptt_tmp.append(i)
                self.ptt_information[2] = '+'.join(self.ptt_tmp)

            # 修正文章名稱冒號問題
            if ': ' in self.ptt_information[2]:
                self.ptt_information[2] = self.ptt_information[2].replace(
                    ': ', '-')

            # 避免標題重複，因此檔案名稱加上作者
            self.ptt_information[2] = self.ptt_information[
                2] + '-' + self.ptt_information[0].split(' ')[0]

            # 過濾掉暱稱
            self.ptt_information[0] = self.ptt_information[0].split(' ')[0]

            # 製作資料夾
            mkdir(self.ptt_information[2])

            # 以 class 找出網址
            urldata = soup.find_all('a', rel="nofollow")
            text_update('開始下載圖片\n')
            for s in urldata:
                if '.jpg' in s.text:
                    self.image_url_count += 1
                    try:
                        self.image_download(s.text)
                    except:
                        text_update('下載圖片過程發生問題\n')
                elif '.png' in s.text:
                    self.image_url_count += 1
                    try:
                        self.image_download(s.text)
                    except:
                        text_update('下載圖片過程發生問題\n')
                elif 'imgur.com/a/' in s.text:
                    try:
                        self.imgur_album_url(s.text, True)
                    except:
                        text_update('下載圖片過程發生問題\n')
                elif 'imgur' in s.text:
                    self.image_url_count += 1
                    try:
                        self.image_download(s.text)
                    except:
                        text_update('下載圖片過程發生問題\n')
                elif 'pixnet' in s.text:
                    pixnet = Pixnet(s.text)
                    pixnet.pixnet_analysis(self.ptt_information[2])
                elif 'https://www.ptt.cc/bbs' in s.text:
                    break

            text_update('圖片下載完畢\n')
            text_update('------------------------------\n')
            text_update('總共 %d 張圖片\n' % self.image_url_count)

            # 抓取文章內容
            txt = soup.find(id='main-content')
            for i in txt.text.split('\n'):
                self.txt_data.append(i)
                self.sentence_count += 1

            # 製作文件
            try:
                self.ptt_word(self.url)
            except Exception as e:
                print(e)
                text_update('下載文章過程發生問題\n')

    # 印出錯誤 number
    def print_info(self, cnt):
        if cnt > 9:
            text_update('第 %d 張圖片下載\n' % cnt)
        else:
            text_update('第  %d 張圖片下載\n' % cnt)

    # 確認imgur網址的格式
    def double_check_imgur(self, checkurl):
        model = 0

        # 如果網址結尾有 JPG 或是 PNG 的話，則可以直接通過確認
        if 'jpg' in checkurl:
            checkurl = re.search('(.*).jpg', checkurl).group(1) + '.jpg'
        elif 'png' in checkurl:
            checkurl = re.search('(.*).png', checkurl).group(1) + '.png'
        elif 'imgur.dcard.tw' in checkurl:
            checkurl = checkurl.replace('imgur.dcard.tw', 'i.imgur.com')
        elif 'i.imgur' not in checkurl:
            imgur = requests.get(checkurl + '.jpg')

            # 確認是否下載成功
            if imgur.status_code == requests.codes.ok:

                # 判斷圖片類型
                typeofpic = imghdr.what(None, imgur.content)
                if typeofpic == None:
                    checkurl = imgur.url
                else:
                    checkurl = imgur.url[:-3] + typeofpic
        # 回傳正確網址
        return checkurl

    # 確認 imgur album 網址下載
    def imgur_album_url(self, checkurl, download=False):
        checkurl += '/layout/blog'
        imgur = requests.get(checkurl)
        imgur_list = []
        # 確認是否下載成功
        if imgur.status_code == requests.codes.ok:
            # 抓出正確的網址
            try:
                image_hash = re.findall('{"hash":"([a-zA-Z0-9]+)".*?"ext":"(\.[a-zA-Z0-9]+)".*?', imgur.text)
            except Exception as e:
                print(e)
            
            for url in image_hash:
                new_url = 'https://i.imgur.com/' + url[0] + url[1]
                if new_url not in imgur_list:
                    imgur_list.append(new_url)
            if download:
                for url in imgur_list:
                    self.image_url_count += 1
                    filename = str(self.image_url_count) + url[-4:]
                    imgcontent = requests.get(url).content
                    try:
                        with open(self.ptt_information[2] + '/' + filename, 'wb') as code:
                            code.write(imgcontent)
                            self.print_info(self.image_url_count)
                        self.format_data.append(url[-4:])
                    except:
                        self.image_url_count -= 1
            else:
                return imgur_list

    # ptt 圖片下載
    def image_download(self, imageurl):
        if 'imgur' in imageurl:
            imageurl = self.double_check_imgur(imageurl)
        else:
            pass

        filename = str(self.image_url_count) + imageurl[-4:]
        imgcontent = requests.get(imageurl).content
        with open(self.ptt_information[2] + '/' + filename, 'wb') as code:
            code.write(imgcontent)
            self.print_info(self.image_url_count)

        try:
            self.format_data.append(imageurl[-4:])
            typeofpic = imghdr.what(self.ptt_information[2] + '/' + filename)
            if typeofpic == None:
                im = Image.open(self.ptt_information[2] + '/' + filename)
                if imageurl[-4:] == '.jpg':
                    im.save(self.ptt_information[2] + '/' + filename, "JPEG")
                elif imageurl[-4:] == '.png':
                    im.save(self.ptt_information[2] + '/' + filename, "PNG")
                im.close()
        except IOError as e:
            text_update('第 %d 張圖片下載失敗\n' % self.image_url_count)
            self.image_url_count -= 1
            print(e)
        except Exception as e:
            text_update('第 %d 張圖片下載失敗\n' % self.image_url_count)
            self.image_url_count -= 1
            print(e)

    # 製作Word文檔
    def ptt_word(self, url):
        num = 1
        txt = docx.Document()
        text_update('------------------------------\n')
        text_update('開始下載文章\n')
        for i in range(1, self.sentence_count):
            if u'文章網址' in self.txt_data[i]:
                break
            if num <= self.image_url_count:
                str1 = self.ptt_information[2] + '/' + str(
                    num) + self.format_data[num - 1]
                if 'imgur.com/a/' in self.txt_data[i]:
                    filter_url = re.findall('[a-zA-z]+://[^\s]*', self.txt_data[i])
                    imgur_list = self.imgur_album_url(filter_url[0])
                    for url in imgur_list:
                        str1 = self.ptt_information[2] + '/' + str(num) + self.format_data[num - 1]
                        try:
                            txt.add_picture(str1, width=Inches(3))
                        except:
                            print("Can't add pic, so add url: ", url)
                            txt.add_paragraph(url)
                        finally:
                            num += 1
                elif 'imgur' in self.txt_data[i]:
                    try:
                        txt.add_picture(str1, width=Inches(3))
                    except:
                        print("Can't add pic, so add url: ", self.txt_data[i])
                        txt.add_paragraph(self.txt_data[i])
                    num += 1
                elif 'jpg' in self.txt_data[i]:
                    try:
                        txt.add_picture(str1, width=Inches(3))
                    except:
                        print("Can't add pic, so add url: ", self.txt_data[i])
                        txt.add_paragraph(self.txt_data[i])
                    num += 1
                elif 'png' in self.txt_data[i]:
                    try:
                        txt.add_picture(str1, width=Inches(3))
                    except:
                        print("Can't add pic, so add url: ", self.txt_data[i])
                        txt.add_paragraph(self.txt_data[i])
                    num += 1
            txt.add_paragraph(self.txt_data[i])

        txt.add_paragraph(u'\n網友' + self.ptt_information[0] + u'分享於' +
                          self.ptt_information[1])
        txt.add_paragraph(url + '\n')
        txt.save(self.ptt_information[2] + '.docx')
        text_update('文章下載完畢\n')
        text_update('------------------------------\n')


'''
下載IG圖片
'''


class Instagram():
    def __init__(self, url):
        self.__url = url
        self.__headers = {
            'user-agent':
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 11_4) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36 Edg/91.0.864.59'
        }
        self.__url_list = []
        self.__video = []
        self.__ins_title = ''
        self.__text = ''
        self.__image_author_fullname = ''

    def analysis_ins(self):
        
        r = requests.get(self.__url + '?__a=1', headers=self.__headers)
        if r.status_code == requests.codes.ok:

            # Get Data
            data = json.loads(r.text)

            # Title
            # self.__ins_title = re.search('https://www.instagram.com/p/(.*)/', self.__url).group(1)
            short_code = data['graphql']['shortcode_media']['shortcode']
            author_name = data['graphql']['shortcode_media']['owner']['username']
            self.__image_author_fullname = data['graphql']['shortcode_media']['owner']['full_name']
            self.__ins_title = short_code + '_' + author_name

            # Create folder
            try:
                mkdir(self.__ins_title)
            except OSError as e:
                print(e)
                print('Filename have some problem, You should change by yourself')
                self.__ins_title = 'tmp_ins_file_name'
                mkdir(self.__ins_title)

            data = json.loads(r.text)
            if 'edge_sidecar_to_children' in data['graphql']['shortcode_media']:
                image_url = data['graphql']['shortcode_media']['edge_sidecar_to_children']['edges']  # data['entry_data']['PostPage'][0]['graphql']['shortcode_media']
                for image in image_url:
                    if image['node']['is_video'] == True:
                        self.__url_list.append(image['node']['video_url'])
                        self.__video.append(True)
                    else:
                        self.__url_list.append(image['node']['display_url'])
                        self.__video.append(False)
            else:
                image = data['graphql']['shortcode_media']
                if image['is_video'] == True:
                    self.__url_list.append(image['video_url'])
                    self.__video.append(True)
                else:
                    self.__url_list.append(image['display_url'])
                    self.__video.append(False)
            
            # Download Image or Video
            text_update('開始下載圖片和影片\n')
            for idx in range(len(self.__url_list)):
                filename = self.__ins_title + '/'
                filename += str(idx + 1) + '.mp4' if self.__video[idx] == True else str(idx + 1) + '.jpg'
                pic = '影片' if self.__video[idx] == True else '圖片'
                try:
                    imgcontent = requests.get(self.__url_list[idx]).content
                    with open(filename, 'wb') as code:
                        code.write(imgcontent)
                    text_update('第 %d 張%s下載\n' % (idx + 1, pic))
                except Exception as e:
                    text_update('第 %d 張%s下載失敗\n' % (idx + 1, pic))
                    print(e)
            text_update('下載完畢\n')
            text_update('------------------------------\n')

            self.__text = data['graphql']['shortcode_media']['edge_media_to_caption']['edges'][0]['node']['text']
            self.ins_txt_download(self.__url)

    def ins_txt_download(self, url):
        text_update('正在下載文章\n')
        filename = self.__ins_title + '.txt'
        with open(self.__ins_title + '/' + filename, 'a') as code:
            code.write(self.__text)
            code.write('\n')
            code.write('文章網址: ' + url)
            code.write('\n')
            code.write('圖片上傳者: ' + self.__image_author_fullname)
            code.write('\n')
        text_update('文章儲存完成\n')
        text_update('------------------------------\n')

'''
下載 Pixnet 文章以及圖片
'''


class Pixnet():
    def __init__(self, url):
        self.__url = url
        self.pixnet_text = []
        self.pixnet_image_url_count = 0
        self.pixnet_sentence_count = 0
        self.pixnet_title = []
        self.pixnet_dir = []
        self.pixnet_headers = {
            'user-agent':
            'Mozilla/5.0 (Macintosh Intel Mac OS X 10_12_6) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/69.0.3497.100 Safari/537.36'
        }

    # 分析 pixnet 文章網址
    def pixnet_analysis(self, pixnetdir=''):
        r = requests.get(self.__url, headers=self.pixnet_headers)

        # 先 encode成 utf-8 的格式
        r.encoding = 'utf-8'

        if r.status_code == requests.codes.ok:
            soup = BeautifulSoup(r.text, 'html.parser')

            # 抓取文章標題並且修正
            title = soup.select('.title a')
            self.pixnet_title = title[0].contents[0] + '(Pixnet)'
            # print(title[0].contents[0].encode('utf-8', 'ignore'))

            if '/' in self.pixnet_title:
                tmp = []
                for i in self.pixnet_title.split('/'):
                    tmp.append(i)
                self.pixnet_title = '_'.join(tmp)

            if pixnetdir != '':
                self.pixnet_dir = pixnetdir + '/' + self.pixnet_title
            else:
                self.pixnet_dir = self.pixnet_title

            # 創建資料夾
            mkdir(self.pixnet_dir)

            # 抓取文章內容及圖片網址
            text = soup.select('div.article-content > div > p')
            text_update('正在下載圖片\n')
            for a in text:
                model = 0
                txt = a.select('img[src]')
                for b in txt:
                    self.pixnet_image_url_count += 1
                    try:
                        self.pixnet_image_download(b.get('src'))
                    except:
                        text_update('下載圖片過程發生問題\n')
                    self.pixnet_text.append(b.get('src'))
                    self.pixnet_sentence_count += 1
                    model += 1
                if model == 0:
                    self.pixnet_text.append(a.text)
                    self.pixnet_sentence_count += 1
            if self.pixnet_image_url_count > 9:
                text_update('總共 %d 張圖片\n' % self.pixnet_image_url_count)
            else:
                text_update('總共  %d 張圖片\n' % self.pixnet_image_url_count)
            text_update('圖片下載完成\n')

            # 下載 pixnet 文章
            try:
                self.pixnet_txt_download(self.__url)
            except:
                text_update('下載文章過程發生問題\n')

    # 下載圖片
    def pixnet_image_download(self, pixnetimageurl):

        if 'https:' not in pixnetimageurl:
            pixnetimageurl = 'https:' + pixnetimageurl
        filename = str(self.pixnet_image_url_count) + pixnetimageurl[-4:]
        imgcontent = requests.get(
            pixnetimageurl, headers=self.pixnet_headers).content
        with open(self.pixnet_dir + '/' + filename, 'wb') as code:
            code.write(imgcontent)
            if self.pixnet_image_url_count > 9:
                text_update('第 %d 張圖片下載\n' % self.pixnet_image_url_count)
            else:
                text_update('第  %d 張圖片下載\n' % self.pixnet_image_url_count)

    #下載 pixnet 文章
    def pixnet_txt_download(self, url):
        text_update('------------------------------\n')
        text_update('正在下載文章\n')
        filename = self.pixnet_title + '.txt'
        with open(self.pixnet_dir + '/' + filename, 'a') as code:
            for i in range(self.pixnet_sentence_count):
                code.write(self.pixnet_text[i])
                code.write('\n')
            code.write('\n')
            code.write("文章網址 :" + url)
        text_update('文章儲存完成\n')
        text_update('------------------------------\n')

# 檢查 titlename 有無超過 U+FFFF
def check_title_name(titlename):
    fix_titlename = ""
    for i in range(len(titlename)):
        if ord(titlename[i]) not in range(65536):
            fix_titlename += '?'
        else:
            fix_titlename += titlename[i]
    return fix_titlename

# 建立資料夾
def mkdir(titlename):
    # 用 path = os.getcwd() 來取得腳本位置
    path = os.getcwd()

    # 創建文章標題的資料夾儲存圖片
    newpath = path + '/' + titlename
    if not os.path.isdir(newpath):
        os.mkdir(newpath)
        text_update('------------------------------\n')
        text_update('文章標題: %s' % check_title_name(titlename))
        text_update('\n------------------------------\n')


state = 0

if __name__ == "__main__":

    windows = tk.Tk()

    # 標題
    windows.title('CYYDownloader')

    # 視窗大小
    windows.geometry('800x500')

    # 設置封面圖片
    img_open = Image.open('title.png')
    img_png = ImageTk.PhotoImage(img_open)
    label_img = tk.Label(windows, image=img_png)
    label_img.pack()

    def clear_box(event):
        web_url.delete(0, tk.END)
        return

    tk.Label(windows, text='下載網址: ', font=('Arial', '20')).place(x=10, y=100)
    get_web_url = tk.StringVar()
    get_web_url.set('請輸入文章網址:')
    web_url = tk.Entry(windows, textvariable=get_web_url, width=70)
    web_url.place(x=100, y=100)
    web_url.bind("<Button>", clear_box)

    # 輸出文字介面
    tk.Label(windows, text='下載狀態: ', font=('Arial', '20')).place(x=10, y=140)
    textbox = tk.Text(windows, height=10, width=80)
    textbox.place(x=10, y=180)

    # 顯示目前狀態
    def text_update(word):
        textbox.insert('insert', word)
        textbox.update()
        textbox.see(tk.END)

    # 彈出錯誤視窗
    def error_message():
        messagebox.showerror(title='Warning', message='輸入無效網址，請重新輸入。')

    # 彈出下載完成
    def finish_download():
        messagebox.showerror(title='Finish!!!!', message='下載完成！！！！！！！')

    # 主程式
    def run():
        global state
        if state == 1:
            textbox.delete(1.0, tk.END)
        url = get_web_url.get()
        if 'ptt' in url:
            state = 1
            ptt = Ptt(url)
            ptt.analysis()
            finish_download()
        elif 'dcard' in url:
            state = 1
            dcard = Dcard(url)
            dcard.dcard_analysis()
            finish_download()
        elif 'youtube' in url:
            state = 1
            yoube = Youtube(url)
            yoube.download_video()
            finish_download()
        elif 'instagram' in url:
            state = 1
            ig = Instagram(url)
            ig.analysis_ins()
            finish_download()
        elif 'pixnet' in url:
            state = 1
            pixnet = Pixnet(url)
            pixnet.pixnet_analysis()
            finish_download()
        elif url == 'q':
            windows.destroy()
        else:
            state = 0
            error_message()

    # 按鈕
    start_button = tk.Button(
        windows,
        text='開始下載',
        width=20,
        height=2,
        font=('Arial', '20'),
        command=run).place(
            x=10, y=360)
    end_button = tk.Button(
        windows,
        text='結束程式',
        width=20,
        height=2,
        font=('Arial', '20'),
        command=windows.destroy).place(
            x=300, y=360)

    # 主循環
    windows.mainloop()
